# -*- coding: utf-8 -*-
"""
会议纪要生成工具
根据会议转写文稿自动生成会议纪要
"""

import sys
import os
libs_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "libs")
if libs_path not in sys.path:
    sys.path.insert(0, libs_path)

# ==================== API 配置区域 ====================
# 请在此处填写您的 API 配置

API_KEY = "f441b58b15bf473099cc7614e24696d6.4MgeIeAHNLwyiDEk"  # 填入您的 API Key
BASE_URL = "https://open.bigmodel.cn/api/paas/v4"  # 填入您的 API Base URL
MODEL_NAME = "glm-5"  # 可根据需要修改模型名称

MAX_TEXT_LENGTH = 50000

# ====================================================

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import json
import threading
import time
from datetime import datetime
from typing import List, Dict, Optional
import copy

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI, RateLimitError
import httpx
import re

DEFAULT_RULES = """你是一个专业的会议纪要助手。请根据会议转写文稿生成会议纪要。
要求：
1. 提取会议基本信息（主题、时间、人员）。
2. 总结会议核心内容，按议题分组，条理清晰。
3. 提取遗留事项，明确责任人和时间。
4. 语言专业、简洁。

【重要】输出要求：必须严格输出 JSON 格式，不包含 Markdown 标记。
【重要】必须严格使用以下 JSON 结构和 Key 名称，不要擅自修改字段名：

{
  "会议信息": {
    "会议主题": "会议主题内容",
    "会议时间": "会议时间",
    "参会人员": "参会人员名单"
  },
  "会议内容": [
    {
      "议题": "议题名称",
      "内容摘要": "该议题的详细内容摘要，可以包含多个要点"
    }
  ],
  "遗留事项": [
    {
      "序号": "1",
      "事项": "具体任务内容",
      "责任人": "负责人姓名",
      "截止时间": "截止日期"
    }
  ]
}

请务必严格使用上述 JSON 结构和 Key 名称：
- 必须使用"会议信息"而不是"会议基本信息"
- 必须使用"会议内容"数组格式，每个议题包含"议题"和"内容摘要"字段
- 必须使用"事项"而不是"事项描述"
- 必须使用"截止时间"而不是"完成时间"
不要修改任何字段名。"""


class MeetingMinutesApp:
    def __init__(self, root):
        self.root = root
        self.root.title("会议纪要生成工具")
        self.root.geometry("700x780")
        self.root.resizable(True, True)
        
        self.minutes_files: List[str] = []
        self.minutes_template_path: str = ""
        self.minutes_rules: str = DEFAULT_RULES
        self.is_processing = False
        self.process_mode = tk.StringVar(value="batch")
        
        self._load_config()
        self._create_widgets()
    
    def _load_config(self):
        config_path = self._get_config_path()
        if os.path.exists(config_path):
            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
                    self.minutes_rules = config.get("rules", DEFAULT_RULES)
                    self.minutes_template_path = config.get("template_path", "")
            except:
                pass
    
    def _save_config(self):
        config_path = self._get_config_path()
        try:
            config = {
                "rules": self.minutes_rules,
                "template_path": self.minutes_template_path
            }
            with open(config_path, "w", encoding="utf-8") as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            messagebox.showerror("保存失败", f"保存配置失败: {str(e)}")
    
    def _get_config_path(self) -> str:
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")
    
    def _create_widgets(self):
        main_frame = tk.Frame(self.root, padx=15, pady=15)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        file_frame = tk.LabelFrame(main_frame, text="文件上传区", padx=10, pady=10)
        file_frame.pack(fill=tk.X, pady=(0, 5))
        
        file_btn_frame = tk.Frame(file_frame)
        file_btn_frame.pack(fill=tk.X)
        
        self.minutes_file_btn = tk.Button(
            file_btn_frame, 
            text="选择转写文稿", 
            command=self._select_minutes_files,
            width=15,
            height=1
        )
        self.minutes_file_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        self.minutes_file_label = tk.Label(file_btn_frame, text="支持格式: DOCX（会议转写文稿）", anchor="w")
        self.minutes_file_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        agenda_frame = tk.LabelFrame(main_frame, text="预设议题（可选）", padx=10, pady=10)
        agenda_frame.pack(fill=tk.X, pady=(0, 5))
        
        agenda_hint = tk.Label(agenda_frame, text="请输入会议议题，每行一个议题。留空则由AI自动分析。", anchor="w", fg="gray")
        agenda_hint.pack(fill=tk.X, pady=(0, 5))
        
        self.agenda_text = scrolledtext.ScrolledText(agenda_frame, wrap=tk.WORD, height=3)
        self.agenda_text.pack(fill=tk.X)
        
        mode_frame = tk.LabelFrame(main_frame, text="文件处理模式", padx=10, pady=10)
        mode_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.mode_batch_rb = tk.Radiobutton(
            mode_frame, 
            text="视为不同会议（批量处理）- 每个文件单独生成一份纪要", 
            variable=self.process_mode, 
            value="batch",
            command=self._on_mode_change
        )
        self.mode_batch_rb.pack(anchor="w")
        
        self.mode_merge_rb = tk.Radiobutton(
            mode_frame, 
            text="视为同一场会议（合并处理）- 所有文件合并生成一份纪要", 
            variable=self.process_mode, 
            value="merge",
            command=self._on_mode_change
        )
        self.mode_merge_rb.pack(anchor="w")
        
        rules_frame = tk.LabelFrame(main_frame, text="规则配置区", padx=10, pady=10)
        rules_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.rules_text = scrolledtext.ScrolledText(rules_frame, wrap=tk.WORD, height=5)
        self.rules_text.pack(fill=tk.X, pady=(0, 5))
        self.rules_text.insert(tk.END, self.minutes_rules)
        
        self.save_rules_btn = tk.Button(
            rules_frame, 
            text="保存规则", 
            command=self._save_rules,
            width=15
        )
        self.save_rules_btn.pack(anchor="w")
        
        template_frame = tk.LabelFrame(main_frame, text="模板设置区", padx=10, pady=10)
        template_frame.pack(fill=tk.X, pady=(0, 10))
        
        template_btn_frame = tk.Frame(template_frame)
        template_btn_frame.pack(fill=tk.X)
        
        self.template_btn = tk.Button(
            template_btn_frame, 
            text="选择模板", 
            command=self._select_template,
            width=15,
            height=1
        )
        self.template_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        self.template_label = tk.Label(template_btn_frame, text="未选择模板", anchor="w")
        self.template_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        if self.minutes_template_path and os.path.exists(self.minutes_template_path):
            self.template_label.config(text=self.minutes_template_path)
        
        action_frame = tk.Frame(main_frame)
        action_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.minutes_process_btn = tk.Button(
            action_frame, 
            text="开始生成会议纪要", 
            command=self._start_minutes_generation,
            width=20,
            height=2,
            state=tk.DISABLED
        )
        self.minutes_process_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        self.minutes_status_label = tk.Label(action_frame, text="请选择转写文稿和模板", anchor="w")
        self.minutes_status_label.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        log_frame = tk.LabelFrame(main_frame, text="处理日志", padx=10, pady=10)
        log_frame.pack(fill=tk.BOTH, expand=True)
        
        self.minutes_progress = scrolledtext.ScrolledText(
            log_frame, 
            wrap=tk.WORD, 
            height=10,
            state=tk.DISABLED
        )
        self.minutes_progress.pack(fill=tk.BOTH, expand=True)
    
    def _on_mode_change(self):
        mode = self.process_mode.get()
        if mode == "merge":
            self._log("已切换为合并模式：所有文件将合并为一场会议处理")
        else:
            self._log("已切换为批量模式：每个文件将单独生成会议纪要")
    
    def _log(self, message: str):
        self.minutes_progress.config(state=tk.NORMAL)
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.minutes_progress.insert(tk.END, f"[{timestamp}] {message}\n")
        self.minutes_progress.see(tk.END)
        self.minutes_progress.config(state=tk.DISABLED)
        self.root.update_idletasks()
    
    def _select_minutes_files(self):
        file_paths = filedialog.askopenfilenames(
            title="选择会议转写文稿",
            filetypes=[
                ("Word 文档", "*.docx"),
                ("所有文件", "*.*")
            ]
        )
        
        if file_paths:
            self.minutes_files = list(file_paths)
            count = len(self.minutes_files)
            self.minutes_file_label.config(text=f"已选择 {count} 个文件")
            self._update_minutes_button_state()
            self._log(f"已选择 {count} 个转写文稿")
            for f in self.minutes_files:
                self._log(f"  - {os.path.basename(f)}")
    
    def _select_template(self):
        file_path = filedialog.askopenfilename(
            title="选择会议纪要模板",
            filetypes=[
                ("Word 文档", "*.docx"),
                ("所有文件", "*.*")
            ]
        )
        
        if file_path:
            self.minutes_template_path = file_path
            self.template_label.config(text=file_path)
            self._save_config()
            self._update_minutes_button_state()
            self._log(f"已选择模板: {os.path.basename(file_path)}")
    
    def _save_rules(self):
        self.minutes_rules = self.rules_text.get("1.0", tk.END).strip()
        self._save_config()
        messagebox.showinfo("保存成功", "规则已保存到 config.json")
    
    def _update_minutes_button_state(self):
        if self.minutes_files and self.minutes_template_path:
            self.minutes_process_btn.config(state=tk.NORMAL)
            self.minutes_status_label.config(text="准备就绪，点击开始生成")
        elif self.minutes_files:
            self.minutes_process_btn.config(state=tk.DISABLED)
            self.minutes_status_label.config(text="请选择模板文件")
        elif self.minutes_template_path:
            self.minutes_process_btn.config(state=tk.DISABLED)
            self.minutes_status_label.config(text="请选择转写文稿")
        else:
            self.minutes_process_btn.config(state=tk.DISABLED)
            self.minutes_status_label.config(text="请选择转写文稿和模板")
    
    def _read_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
    
    def _get_preset_agendas(self) -> List[str]:
        agenda_content = self.agenda_text.get("1.0", tk.END).strip()
        if not agenda_content:
            return []
        
        agendas = []
        for line in agenda_content.split("\n"):
            line = line.strip()
            if line:
                agendas.append(line)
        
        return agendas
    
    def _generate_minutes_with_llm(self, text: str) -> Optional[Dict]:
        if not API_KEY or not BASE_URL:
            self._log("错误: 请先配置 API_KEY 和 BASE_URL")
            return None
        
        preset_agendas = self._get_preset_agendas()
        
        agenda_instruction = ""
        if preset_agendas:
            agenda_list = "、".join(preset_agendas)
            self._log(f"  检测到预设议题: {agenda_list}，将引导AI按此结构生成")
            agenda_instruction = f"""

【重要】请优先根据以下预设议题来组织会议纪要的"会议内容"部分。请将会议内容归类到这些议题下。
预设议题列表：{json.dumps(preset_agendas, ensure_ascii=False)}
如果会议中有超出预设议题的内容，请归类到"其他事项"中。"""
        
        prompt = f"""{self.minutes_rules}{agenda_instruction}

会议转写文稿：
{text}

请严格按以下 JSON 格式输出，不要包含任何 Markdown 标记：
{{
  "会议信息": {{
    "会议主题": "...",
    "会议时间": "...",
    "参会人员": "..."
  }},
  "会议内容": [
    {{
      "议题": "议题名称",
      "内容摘要": "该议题的详细内容摘要，可以包含多个要点"
    }}
  ],
  "遗留事项": [
    {{"序号": "1", "事项": "...", "责任人": "...", "截止时间": "..."}},
    {{"序号": "2", "事项": "...", "责任人": "...", "截止时间": "..."}}
  ]
}}"""
        
        start_time = time.time()
        
        retry_wait_times = [10, 30, 60]
        max_retries = len(retry_wait_times)
        retry_count = 0
        
        http_client = httpx.Client(timeout=120.0)
        client = OpenAI(
            api_key=API_KEY, 
            base_url=BASE_URL,
            http_client=http_client
        )
        
        while True:
            try:
                self._log(f"  正在发送请求，模型版本: {MODEL_NAME}，文本长度: {len(text)} 字符")
                self._log("  正在调用大模型 API...")
                
                response = client.chat.completions.create(
                    model=MODEL_NAME,
                    messages=[
                        {"role": "system", "content": "你是一个专业的会议纪要助手，必须严格输出 JSON 格式数据。"},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3
                )
                
                elapsed_time = time.time() - start_time
                self._log(f"  API 调用耗时: {elapsed_time:.2f} 秒")
                
                result_text = response.choices[0].message.content.strip()
                self._log(f"  大模型返回内容 (前300字符): {result_text[:300]}...")
                
                json_start = result_text.find("{")
                json_end = result_text.rfind("}") + 1
                if json_start != -1 and json_end > json_start:
                    json_str = result_text[json_start:json_end]
                    result = json.loads(json_str)
                    self._log("  JSON 解析成功")
                    return result
                else:
                    self._log("  无法从返回结果中解析 JSON")
                    return None
            
            except RateLimitError as e:
                retry_count += 1
                if retry_count <= max_retries:
                    wait_time = retry_wait_times[retry_count - 1]
                    self._log(f"  检测到服务繁忙，正在自动排队等待（第{retry_count}次尝试），{wait_time}秒后重试...")
                    time.sleep(wait_time)
                else:
                    elapsed_time = time.time() - start_time
                    self._log(f"  已重试{max_retries}次，仍然失败（总耗时: {elapsed_time:.1f}秒）")
                    self.root.after(0, lambda: messagebox.showerror("服务繁忙", "服务器限流，已重试3次仍然失败，请稍后再试"))
                    return None
            
            except httpx.TimeoutException:
                elapsed_time = time.time() - start_time
                self._log(f"  请求超时（已等待 {elapsed_time:.1f} 秒），请检查网络连接")
                self.root.after(0, lambda: messagebox.showerror("超时错误", "服务器响应超时，请检查网络连接或稍后重试"))
                return None
            
            except httpx.ConnectError:
                self._log("  网络连接失败，请检查是否需要代理或网络是否通畅")
                self.root.after(0, lambda: messagebox.showerror("网络错误", "网络连接失败，请检查是否需要代理或网络是否通畅"))
                return None
            
            except httpx.HTTPStatusError as e:
                status_code = e.response.status_code
                error_detail = e.response.text[:500]
                self._log(f"  HTTP 错误: 状态码 {status_code}")
                self._log(f"  错误详情: {error_detail}")
                self.root.after(0, lambda msg=f"HTTP 错误 {status_code}\n\n{error_detail[:200]}": messagebox.showerror("HTTP 错误", msg))
                return None
            
            except json.JSONDecodeError as e:
                elapsed_time = time.time() - start_time
                error_msg = str(e)
                self._log(f"  JSON 解析失败: {error_msg}（耗时: {elapsed_time:.2f} 秒）")
                return None
            
            except Exception as e:
                elapsed_time = time.time() - start_time
                error_type = type(e).__name__
                error_msg = str(e)
                self._log(f"  API 调用失败 [{error_type}]: {error_msg}（耗时: {elapsed_time:.2f} 秒）")
                self.root.after(0, lambda msg=f"{error_type}: {error_msg}": messagebox.showerror("API 错误", f"API 调用失败\n\n{msg}"))
                return None
    
    def _replace_cell_text_preserve_style(self, cell, new_text: str):
        if cell.paragraphs:
            para = cell.paragraphs[0]
            if para.runs:
                first_run = para.runs[0]
                font_name = first_run.font.name
                font_size = first_run.font.size
                font_bold = first_run.font.bold
                font_color = first_run.font.color.rgb if first_run.font.color.rgb else None
                
                para.clear()
                new_run = para.add_run(new_text)
                if font_name:
                    new_run.font.name = font_name
                if font_size:
                    new_run.font.size = font_size
                if font_bold is not None:
                    new_run.font.bold = font_bold
                if font_color:
                    new_run.font.color.rgb = font_color
            else:
                para.clear()
                para.add_run(new_text)
        else:
            cell.text = new_text
    
    def _fill_template_table(self, doc, data: Dict):
        meeting_info = data.get("会议信息") or data.get("会议基本信息") or {}
        self._log(f"    会议信息字段: {list(meeting_info.keys())}")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    for key, value in meeting_info.items():
                        placeholder = "{{" + key + "}}"
                        if placeholder in cell_text:
                            new_text = cell_text.replace(placeholder, str(value) if value else "待定")
                            self._replace_cell_text_preserve_style(cell, new_text)
                            self._log(f"    替换表格占位符: {placeholder} -> {value}")
    
    def _set_paragraph_spacing(self, paragraph, before_pt=6, after_pt=6, line_spacing=1.5):
        pf = paragraph.paragraph_format
        pf.space_before = Pt(before_pt)
        pf.space_after = Pt(after_pt)
        pf.line_spacing = line_spacing
    
    def _set_run_font(self, run, font_size=12, bold=False, font_name="宋体"):
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = font_name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.insert(0, rFonts)
    
    def _add_indent(self, paragraph, first_line_chars=2):
        pf = paragraph.paragraph_format
        pf.first_line_indent = Pt(24)
    
    def _split_content_into_paragraphs(self, content: str) -> List[str]:
        paragraphs = []
        lines = re.split(r'\n+', content)
        for line in lines:
            line = line.strip()
            if not line:
                continue
            sub_parts = re.split(r'(?=\d+[\.、）])', line)
            if len(sub_parts) > 1:
                for part in sub_parts:
                    part = part.strip()
                    if part:
                        paragraphs.append(part)
            else:
                paragraphs.append(line)
        return paragraphs
    
    def _fill_template_paragraph(self, doc, content):
        if content is None:
            content = ""
        
        placeholder_para = None
        placeholder_idx = -1
        
        for idx, para in enumerate(doc.paragraphs):
            if "{{会议内容}}" in para.text:
                placeholder_para = para
                placeholder_idx = idx
                break
        
        if placeholder_para is None:
            self._log("    未找到 {{会议内容}} 占位符")
            return False
        
        parent = placeholder_para._element.getparent()
        placeholder_element = placeholder_para._element
        
        if isinstance(content, list):
            self._log(f"    会议内容为数组格式，共 {len(content)} 个议题，开始分级排版...")
            
            for i, item in enumerate(content):
                if isinstance(item, dict):
                    title = item.get("议题") or item.get("标题") or item.get("议题名称") or ""
                    summary = item.get("内容") or item.get("摘要") or item.get("内容摘要") or item.get("详情") or ""
                else:
                    title = str(item)
                    summary = ""
                
                if not title and not summary:
                    continue
                
                title_para = doc.add_paragraph()
                title_text = f"议题{i+1}：{title}" if title else f"议题{i+1}"
                title_run = title_para.add_run(title_text)
                self._set_run_font(title_run, font_size=14, bold=True, font_name="黑体")
                self._set_paragraph_spacing(title_para, before_pt=12, after_pt=6, line_spacing=1.5)
                parent.insert(parent.index(placeholder_element), title_para._element)
                self._log(f"      插入议题标题: {title_text[:30]}...")
                
                if summary:
                    sub_paras = self._split_content_into_paragraphs(summary)
                    for sub_content in sub_paras:
                        content_para = doc.add_paragraph()
                        content_run = content_para.add_run(sub_content)
                        self._set_run_font(content_run, font_size=12, bold=False, font_name="宋体")
                        self._set_paragraph_spacing(content_para, before_pt=0, after_pt=6, line_spacing=1.5)
                        self._add_indent(content_para, first_line_chars=2)
                        parent.insert(parent.index(placeholder_element), content_para._element)
        else:
            self._log("    会议内容为文本格式，开始智能分段排版...")
            content = str(content)
            paragraphs = self._split_content_into_paragraphs(content)
            
            for para_text in paragraphs:
                new_para = doc.add_paragraph()
                run = new_para.add_run(para_text)
                self._set_run_font(run, font_size=12, bold=False, font_name="宋体")
                self._set_paragraph_spacing(new_para, before_pt=0, after_pt=6, line_spacing=1.5)
                self._add_indent(new_para, first_line_chars=2)
                parent.insert(parent.index(placeholder_element), new_para._element)
        
        parent.remove(placeholder_element)
        self._log("    已删除占位符段落，排版完成")
        return True
    
    def _fill_template_todos_table(self, doc, todos: List[Dict]):
        if not todos:
            self._log("    遗留事项列表为空")
            return
        
        self._log(f"    遗留事项数量: {len(todos)}")
        self._log(f"    第一条数据字段名: {list(todos[0].keys())}")
        self._log(f"    第一条数据内容: {todos[0]}")
        
        todo_table = None
        template_row_idx = -1
        
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                row_text = ""
                for cell in row.cells:
                    row_text += cell.text
                
                if "{{序号}}" in row_text or "{{事项}}" in row_text:
                    todo_table = table
                    template_row_idx = row_idx
                    break
            
            if todo_table:
                break
        
        if not todo_table or template_row_idx < 0:
            self._log("    未找到遗留事项模板行")
            return
        
        self._log(f"    找到遗留事项表格，模板行索引: {template_row_idx}")
        
        template_row = todo_table.rows[template_row_idx]
        
        for i, todo in enumerate(todos):
            new_row = copy.deepcopy(template_row._element)
            todo_table._tbl.append(new_row)
            
            new_row_obj = todo_table.rows[-1]
            
            seq = todo.get("序号") or todo.get("编号") or todo.get("序") or str(i + 1)
            seq_field = "序号" if todo.get("序号") else ("编号" if todo.get("编号") else ("序" if todo.get("序") else "默认序号"))
            
            item = todo.get("事项") or todo.get("事项描述") or todo.get("任务") or todo.get("内容") or todo.get("待办") or ""
            item_field = "事项" if todo.get("事项") else ("事项描述" if todo.get("事项描述") else ("任务" if todo.get("任务") else ("内容" if todo.get("内容") else "待办")))
            
            person = todo.get("责任人") or todo.get("负责人") or todo.get("执行人") or todo.get("责任") or ""
            person_field = "责任人" if todo.get("责任人") else ("负责人" if todo.get("负责人") else ("执行人" if todo.get("执行人") else "责任"))
            
            deadline = todo.get("截止时间") or todo.get("完成时间") or todo.get("期限") or todo.get("时限") or ""
            deadline_field = "截止时间" if todo.get("截止时间") else ("完成时间" if todo.get("完成时间") else ("期限" if todo.get("期限") else "时限"))
            
            if i == 0:
                self._log(f"    字段映射: 序号={seq_field}, 事项={item_field}, 责任人={person_field}, 截止时间={deadline_field}")
            
            seq = str(seq) if seq else str(i + 1)
            item = str(item) if item else "待定"
            person = str(person) if person else "待定"
            deadline = str(deadline) if deadline else "待定"
            
            for cell in new_row_obj.cells:
                cell_text = cell.text
                new_text = cell_text.replace("{{序号}}", seq)
                new_text = new_text.replace("{{事项}}", item)
                new_text = new_text.replace("{{责任人}}", person)
                new_text = new_text.replace("{{截止时间}}", deadline)
                self._replace_cell_text_preserve_style(cell, new_text)
            
            self._log(f"    添加遗留事项行 {i + 1}: {item[:30]}...")
        
        template_row._element.getparent().remove(template_row._element)
        self._log("    已删除模板行")
    
    def _generate_minutes_from_template(self, source_file: str, data: Dict, output_name_suffix: str = "") -> Optional[str]:
        try:
            self._log("  正在加载模板文件...")
            doc = Document(self.minutes_template_path)
            
            self._log(f"  JSON 数据结构: {list(data.keys())}")
            
            meeting_info = data.get("会议信息") or data.get("会议基本信息") or {}
            self._log(f"  会议信息: {meeting_info}")
            
            content = data.get("会议内容") or data.get("会议核心议题") or ""
            if isinstance(content, list):
                self._log(f"  会议内容为数组格式，共 {len(content)} 条")
            else:
                self._log(f"  会议内容长度: {len(str(content))} 字符")
            
            todos = data.get("遗留事项") or data.get("待办事项") or []
            self._log(f"  遗留事项数量: {len(todos)} 条")
            
            self._log("  正在填充会议信息表格...")
            self._fill_template_table(doc, data)
            
            self._log("  正在填充会议内容段落...")
            self._fill_template_paragraph(doc, content)
            
            if todos:
                self._log(f"  正在填充遗留事项表格 ({len(todos)} 条)...")
                self._fill_template_todos_table(doc, todos)
            
            source_dir = os.path.dirname(source_file)
            base_name = os.path.splitext(os.path.basename(source_file))[0]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            if output_name_suffix:
                filename = f"会议纪要_{output_name_suffix}_{timestamp}.docx"
            else:
                filename = f"会议纪要_{base_name}_{timestamp}.docx"
            filepath = os.path.join(source_dir, filename)
            
            self._log(f"  正在保存文件: {filepath}")
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            error_msg = str(e)
            self._log(f"  生成会议纪要失败: {error_msg}")
            self.root.after(0, lambda msg=error_msg: messagebox.showerror("错误", f"生成会议纪要失败: {msg}"))
            return None
    
    def _process_batch_mode(self) -> tuple:
        success_count = 0
        output_files = []
        
        for i, file_path in enumerate(self.minutes_files, 1):
            self._log(f"\n{'='*50}")
            self._log(f"正在处理第 {i}/{len(self.minutes_files)} 个文件: {os.path.basename(file_path)}")
            
            try:
                self._log("  正在读取转写文稿...")
                text = self._read_docx(file_path)
                if not text:
                    self._log("  读取文件失败")
                    continue
                
                self._log(f"  文稿读取完成，文本长度: {len(text)} 字符")
                
                self._log("  正在调用大模型生成会议纪要数据...")
                minutes_data = self._generate_minutes_with_llm(text)
                
                if not minutes_data:
                    self._log("  大模型调用失败")
                    continue
                
                output_path = self._generate_minutes_from_template(file_path, minutes_data)
                
                if output_path:
                    success_count += 1
                    output_files.append(output_path)
                    self._log(f"  会议纪要生成成功: {output_path}")
                else:
                    self._log("  会议纪要生成失败")
                    
            except Exception as e:
                self._log(f"  处理文件出错: {str(e)}")
        
        return success_count, output_files
    
    def _process_merge_mode(self) -> tuple:
        self._log(f"\n{'='*50}")
        self._log("合并模式：正在读取所有文件...")
        
        all_texts = []
        for i, file_path in enumerate(self.minutes_files, 1):
            self._log(f"  读取第 {i}/{len(self.minutes_files)} 个文件: {os.path.basename(file_path)}")
            text = self._read_docx(file_path)
            if text:
                all_texts.append(f"【分段{i}】\n{text}")
        
        if not all_texts:
            self._log("  所有文件读取失败")
            return 0, []
        
        merged_text = "\n\n".join(all_texts)
        self._log(f"  所有文件合并完成，总文本长度: {len(merged_text)} 字符")
        
        if len(merged_text) > MAX_TEXT_LENGTH:
            self._log(f"  警告: 文件内容过长 ({len(merged_text)} 字符)，建议减少文件数量或分开处理")
            self.root.after(0, lambda: messagebox.showwarning(
                "内容过长", 
                f"合并后的文本长度 ({len(merged_text)} 字符) 可能超过模型限制。\n\n建议：\n1. 减少文件数量\n2. 使用批量模式分开处理"
            ))
            return 0, []
        
        self._log("  正在调用大模型生成会议纪要数据...")
        minutes_data = self._generate_minutes_with_llm(merged_text)
        
        if not minutes_data:
            self._log("  大模型调用失败")
            return 0, []
        
        source_dir = os.path.dirname(self.minutes_files[0])
        output_path = self._generate_minutes_from_template(
            self.minutes_files[0], 
            minutes_data, 
            output_name_suffix="合并会议"
        )
        
        if output_path:
            self._log(f"  会议纪要生成成功: {output_path}")
            return 1, [output_path]
        else:
            self._log("  会议纪要生成失败")
            return 0, []
    
    def _start_minutes_generation_thread(self):
        if not API_KEY or not BASE_URL:
            self._log("错误: 请先配置 API_KEY 和 BASE_URL")
            self.root.after(0, lambda: messagebox.showerror("错误", "请先在代码开头配置 API_KEY 和 BASE_URL"))
            return
        
        self.minutes_file_btn.config(state=tk.DISABLED)
        self.template_btn.config(state=tk.DISABLED)
        self.minutes_process_btn.config(state=tk.DISABLED)
        self.is_processing = True
        
        try:
            mode = self.process_mode.get()
            
            if mode == "merge":
                self._log("\n开始处理（合并模式）...")
                success_count, output_files = self._process_merge_mode()
            else:
                self._log("\n开始处理（批量模式）...")
                success_count, output_files = self._process_batch_mode()
        
        except Exception as e:
            self._log(f"处理过程发生错误: {str(e)}")
            success_count = 0
            output_files = []
        
        finally:
            self.root.after(0, lambda: self.minutes_file_btn.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.template_btn.config(state=tk.NORMAL))
            self.root.after(0, lambda: self._update_minutes_button_state())
            self.is_processing = False
        
        if success_count > 0:
            msg = f"处理完成！\n成功生成 {success_count} 份会议纪要"
            if output_files:
                if len(output_files) == 1:
                    msg += f"\n\n文件已保存至:\n{output_files[0]}"
                else:
                    msg += f"\n\n文件已保存至:\n{os.path.dirname(output_files[0])}"
            final_msg = msg
            self.root.after(0, lambda m=final_msg: messagebox.showinfo("完成", m))
        else:
            self.root.after(0, lambda: messagebox.showwarning("提示", "未能生成任何会议纪要"))
    
    def _start_minutes_generation(self):
        if self.is_processing:
            messagebox.showwarning("提示", "正在处理中，请稍候")
            return
        
        if not self.minutes_files:
            messagebox.showwarning("提示", "请先选择转写文稿")
            return
        
        if not self.minutes_template_path:
            messagebox.showwarning("提示", "请先选择模板文件")
            return
        
        thread = threading.Thread(target=self._start_minutes_generation_thread, daemon=True)
        thread.start()


def main():
    root = tk.Tk()
    app = MeetingMinutesApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
